from docx import Document

# Cria um novo documento
doc = Document()

# Adiciona o título
doc.add_heading('Currículo', level=1)

# Adiciona as informações de contato
doc.add_heading('Informações de Contato', level=2)
doc.add_paragraph('Telefone: 41992903603')
doc.add_paragraph('Email: leandrooliveiragoncallos@gmail.com')
doc.add_paragraph('LinkedIn: www.linkedin.com/in/goncallos')
doc.add_paragraph('GitHub: https://github.com/goncallos')

# Adiciona a seção de qualificações
doc.add_heading('Qualificações', level=2)
qualifications = [
    "Análise e Desenvolvimento de Sistemas",
    "Técnico em Fotônica",
    "Técnico em Eletrônica",
    "Técnico em Redes",
    "Técnico em Computadores",
    "Técnico em Fibra Óptica"
]
for qualification in qualifications:
    doc.add_paragraph(qualification)

# Adiciona a seção de experiência profissional
doc.add_heading('Experiência Profissional', level=2)
experiences = [
    "Promotor de Vendas na Zwirtes Produtos Alimentícios LTDA (Outubro de 2022 a Dezembro de 2022)",
    "Técnico Instalador e Reparador de Linhas de Fibra Óptica na Fiber x Service LTDA (Dezembro de 2021 a Agosto de 2022)",
    "Técnico de Fibra Óptica na Fibracem Teleinformática Ltda (Setembro de 2020 a Agosto de 2021)",
    "Técnico em Fotônica na Fibracem Teleinformática Ltda (Setembro de 2020 a Agosto de 2021)",
    "Instalador de Equipamentos de Comutação em Telefonia na Sim Internet e Telecomunicações LTDA (Agosto de 2020 a Outubro de 2020)",
    "Alimentador de Linha de Produção na Sunshine Cortinas e Persianas LTDA (Novembro de 2019 a Fevereiro de 2020)",
    "Técnico de Redes na Eagle Net Tecnologia LTDA (Junho de 2019 a Agosto de 2019)",
    "Técnico em Manutenção de Equipamentos de Informática na Egtech Telecom LTDA (Abril de 2018 a Junho de 2019)",
    "Assistente Pleno em Tecnologia da Informação na Jose Alves Ramos Eletrônicos ME (Junho de 2017 a Março de 2018)",
    "Técnico de Fibra Óptica Sênior na Dominium Instalações e Serviços (Outubro de 2013 a Dezembro de 2015)",
    "Instalador de Linhas Telefônicas (IRLA) na Ericsson Gestão e Serviços de Telecomunicações LTDA (Novembro de 2010 a Agosto de 2011)",
    "Auxiliar Pleno em Tecnologia da Informação na Ability Tecnologia e Serviços SA (Junho de 2008 a Julho de 2010)",
    "Auxiliar de Escritório em Geral na SBF Comércio de Produtos Esportivos S.A. (Dezembro de 2006 a Fevereiro de 2008)",
    "Supervisor de Logística na Pronto Logística Ltda (Agosto de 2005 a Novembro de 2006)",
    "Inspetor de Qualidade na Sommantec Recursos Humanos LTDA (Abril de 2005 a Janeiro de 2006)",
    "Operador de Produção na BR System Technology Indústria e Informática LTDA (Novembro de 2003 a Dezembro de 2004)",
    "Operador de Produção na Forbo Siegling Brasil LTDA (Fevereiro de 2002 a Março de 2003)",
    "Moldador de Plástico por Injeção na Bitron do Brasil Componentes Eletromecânicos LTDA (Junho de 2001 a Setembro de 2001)",
    "Moldador de Plástico por Injeção na Piter Pan Indústria e Comércio LTDA (Outubro de 1994 a Maio de 2001)"
]
for experience in experiences:
    doc.add_paragraph(experience)

# Adiciona a seção de cursos de programação
doc.add_heading('Cursos de Programação', level=2)
programming_courses = [
    "Curso de Desenvolvimento Web com JavaScript - [Nome da Instituição]",
    "Curso de Programação em Python - [Nome da Instituição]",
    "Curso de Desenvolvimento em Java - [Nome da Instituição]",
    "Curso de Banco de Dados MySQL - [Nome da Instituição]",
    "Curso de PHP e Desenvolvimento Backend - [Nome da Instituição]"
]
for course in programming_courses:
    doc.add_paragraph(course)

# Salva o documento
doc.save('curriculo_leandro.docx')
