from docx import Document # type: ignore

# Criando um novo documento Word
doc = Document()

# Título do documento
doc.add_heading('Histórico de Trabalho - Leandro de Oliveira Gonçallo', 0)

# Adicionando as experiências profissionais
experiences = [
    {
        "cargo": "PROMOTOR DE VENDAS",
        "empresa": "ZWIRTES PRODUTOS ALIMENTICIOS LTDA - São José dos Pinhais, PR",
        "periodo": "Outubro de 2022 a Dezembro de 2022",
        "descricao": "Promotor de vendas de produtos alimentícios."
    },
    {
        "cargo": "Técnico instalador e reparador de linhas de fibra ótica",
        "empresa": "Fiber x Service LTDA - Curitiba, PR",
        "periodo": "Dezembro de 2021 a Agosto de 2022",
        "descricao": "Realizando reparo e instalação de link de internet."
    },
    {
        "cargo": "Técnico de Fibra Óptica",
        "empresa": "Fibracem Teleinformática Ltda - Pinhais, PR",
        "periodo": "Setembro de 2020 a Agosto de 2021",
        "descricao": (
            "Instalação de redes de fibra óptica / FTTX/XPOM/POL.\n"
            "Operação de redes de fibra óptica /FTTX/XPOM/POL.\n"
            "Desenvolvimento de redes FTTx, PON, incluindo: FTTN, FTTC, FTTH, FTTB, FTTD, Fiber to the apartament."
        )
    },
    {
        "cargo": "TECNICO EM FOTONICA",
        "empresa": "Fibracem Teleinformática Ltda - Pinhais, PR",
        "periodo": "Setembro de 2020 a Agosto de 2021",
        "descricao": (
            "1. Operação e Manutenção de Equipamentos ópticos.\n"
            "2. Montagem de Dispositivos Ópticos.\n"
            "3. Testes e Medidas.\n"
            "4. Desenvolvimento de Protótipos ópticos.\n"
            "5. Suporte Técnico a sistemas ópticos.\n"
            "6. Calibração de Equipamentos ópticos.\n"
            "7. Documentação de procedimentos e resultados.\n"
            "8. Pesquisa e Desenvolvimento na área de fotônica.\n"
            "9. Treinamento de profissionais em tecnologias ópticas."
        )
    },
    {
        "cargo": "Instalador - Preparador de Equipamentos de Comutação em Telefonia",
        "empresa": "Sim Internet e Telecomunicações LTDA - São José dos Pinhais, PR",
        "periodo": "Agosto de 2020 a Outubro de 2020",
        "descricao": (
            "Instalação de redes de fibra óptica e manutenção de equipamentos de comutação.\n"
            "Atendimento técnico a usuários e suporte em sistemas de informática."
        )
    },
    {
        "cargo": "ALIMENTADOR DE LINHA DE PRODUÇÃO",
        "empresa": "SUNSHINE CORTINAS E PERSIANAS LTDA - São José dos Pinhais, PR",
        "periodo": "Novembro de 2019 a Fevereiro de 2020",
        "descricao": (
            "1. Abastecimento de Materiais na linha de produção.\n"
            "2. Monitoramento da Linha e Reposição de Materiais.\n"
            "3. Inspeção Visual de produtos.\n"
            "4. Manutenção Básica de equipamentos de produção.\n"
            "5. Separação de Produtos acabados.\n"
            "6. Cumprimento de Padrões de Segurança.\n"
            "7. Trabalho em Equipe e Registro de Dados de produção."
        )
    },
    {
        "cargo": "Técnico de Redes (Fibra óptica)",
        "empresa": "EAGLE NET TECNOLOGIA LTDA - São José dos Pinhais, PR",
        "periodo": "Junho de 2019 a Agosto de 2019",
        "descricao": (
            "Instalação de redes de fibra óptica e suporte técnico a serviços de internet e TV.\n"
            "Diagnóstico de problemas e manutenção de rede wireless e equipamentos."
        )
    },
    {
        "cargo": "Técnico em Manutenção de Equipamentos de Informática",
        "empresa": "Egtech Telecom LTDA - Borda do Campo, PR",
        "periodo": "Abril de 2018 a Junho de 2019",
        "descricao": (
            "Instalação de redes de fibra óptica e suporte técnico.\n"
            "Manutenção e configuração de equipamentos e softwares."
        )
    },
    {
        "cargo": "Assistente Pleno / Tecnologia da Informação",
        "empresa": "Jose Alves Ramos Eletronicos ME - São José dos Pinhais, PR",
        "periodo": "Junho de 2017 a Março de 2018",
        "descricao": (
            "Suporte técnico a serviços de internet e TV.\n"
            "Instalação, configuração e manutenção de redes e equipamentos."
        )
    },
    {
        "cargo": "Técnico de Fibra Óptica Sênior",
        "empresa": "Dominium Instalações e Serviços - São Paulo, SP",
        "periodo": "Outubro de 2013 a Dezembro de 2015",
        "descricao": (
            "Execução de instalações de internet e IPTV através de fusões ópticas.\n"
            "Instalação e manutenção de cabos de fibra óptica."
        )
    },
    {
        "cargo": "Instalador de Linhas Telefônicas (IRLA)",
        "empresa": "ERICSON GESTÃO E SERVIÇOS DE TELECOMUNICAÇÕES LTDA - São Paulo, SP",
        "periodo": "Novembro de 2010 a Agosto de 2011",
        "descricao": "Instalação de linhas telefônicas e reparos."
    },
    {
        "cargo": "AUXILIAR DE ESCRITÓRIO EM GERAL",
        "empresa": "SBF COMERCIO DE PRODUTOS ESPORTIVOS S.A. - Lapa, SP",
        "periodo": "Dezembro de 2006 a Fevereiro de 2008",
        "descricao": (
            "1. Atendimento Telefônico e Presencial.\n"
            "2. Arquivamento e Organização de Documentos.\n"
            "3. Entrada de Dados em sistemas.\n"
            "4. Preparação e envio de Correspondência.\n"
            "5. Agendamento e Organização de Reuniões.\n"
            "6. Suporte Administrativo a diversos departamentos.\n"
            "7. Controle de Estoque de Materiais de Escritório.\n"
            "8. Emissão de Notas Fiscais e Documentos Fiscais Simples.\n"
            "9. Auxílio em Atividades Financeiras e de Recursos Humanos.\n"
            "10. Utilização de Softwares de Escritório."
        )
    },
    {
        "cargo": "Supervisor de Logística",
        "empresa": "Pronto Logística Ltda - Itapevi, SP",
        "periodo": "Agosto de 2005 a Novembro de 2006",
        "descricao": (
            "1. Planejamento Logístico e Gestão de Estoques.\n"
            "2. Coordenação de Transporte e Gerenciamento de Equipe.\n"
            "3. Implementação de Sistemas Logísticos e Negociação com Fornecedores.\n"
            "4. Controle de Qualidade e Análise de Desempenho Logístico."
        )
    },
    {
        "cargo": "Inspetor de Qualidade",
        "empresa": "SOMMATEC RECURSOS HUMANOS LTDA - Itapevi, SP",
        "periodo": "Abril de 2005 a Janeiro de 2006",
        "descricao": (
            "1. Inspeção de Produtos e realização de Testes e Medidas.\n"
            "2. Monitoramento de Processos de Produção e Documentação de Resultados.\n"
            "3. Identificação de Não Conformidades e Treinamento em Qualidade."
        )
    },
    {
        "cargo": "Operador de Produção",
        "empresa": "BR SYSTEM TECHNOLOGY INDUSTRIA E INFORMATICA LTDA - Itapevi, SP",
        "periodo": "Novembro de 2003 a Dezembro de 2004",
        "descricao": (
            "1. Operação de Máquinas e Equipamentos de produção.\n"
            "2. Montagem e Desmontagem de Equipamentos.\n"
            "3. Controle de Parâmetros e Inspeção Visual.\n"
            "4. Abastecimento de Insumos e Manutenção Básica."
        )
    },
    {
        "cargo": "Operador de Produção",
        "empresa": "FORBO SIEGLING BRASIL LTDA - Itapevi, SP",
        "periodo": "Fevereiro de 2002 a Março de 2003",
        "descricao": (
            "1. Operação de Máquinas e Equipamentos de produção.\n"
            "2. Montagem e Desmontagem de Equipamentos.\n"
            "3. Controle de Parâmetros e Inspeção Visual.\n"
            "4. Abastecimento de Insumos e Manutenção Básica."
        )
    }
]

# Adicionando as informações ao documento
for experience in experiences:
    doc.add_heading(experience["cargo"], level=1)
    doc.add_paragraph(f'Empresa: {experience["empresa"]}')
    doc.add_paragraph(f'Período: {experience["periodo"]}')
    doc.add_paragraph('Descrição:')
    doc.add_paragraph(experience["descricao"])

# Salvando o documento
doc.save('historico_de_trabalho.docx')
